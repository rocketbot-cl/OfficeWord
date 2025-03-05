# coding: utf-8
"""
Base para desarrollo de modulos externos.
Para obtener el modulo/Funcion que se esta llamando:
     GetParams("module")

Para obtener las variables enviadas desde formulario/comando Rocketbot:
    var = GetParams(variable)
    Las "variable" se define en forms del archivo package.json

Para modificar la variable de Rocketbot:
    SetVar(Variable_Rocketbot, "dato")

Para obtener una variable de Rocketbot:
    var = GetVar(Variable_Rocketbot)

Para obtener la Opcion seleccionada:
    opcion = GetParams("option")


Para instalar librerias se debe ingresar por terminal a la carpeta "libs"
    
    pip install <package> -t .

"""
import os
import sys

# This lines is to linter
# -----------------------------------
GetParams = GetParams #type:ignore
tmp_global_obj = tmp_global_obj #type:ignore
PrintException = PrintException #type:ignore
SetVar = SetVar #type:ignore
GetGlobals = GetGlobals #type:ignore

# Add modules libraries to Rocektbot
# -----------------------------------
base_path = tmp_global_obj["basepath"]
cur_path = os.path.join(base_path, 'modules', 'OfficeWord', 'libs')

cur_path_x64 = os.path.join(cur_path, 'Windows' + os.sep +  'x64' + os.sep)
cur_path_x86 = os.path.join(cur_path, 'Windows' + os.sep +  'x86' + os.sep)

if sys.maxsize > 2**32 and cur_path_x64 not in sys.path:
        sys.path.append(cur_path_x64)
if sys.maxsize <= 2**32 and cur_path_x86 not in sys.path:
        sys.path.append(cur_path_x86)

import traceback

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Mm
    import docx2txt
    from subprocess import Popen, PIPE
    from docx.oxml.shared import qn
    import docx.oxml

    docto = os.path.join(cur_path.replace("libs", "bin"), "docto.exe")

    class DocxModule:

        def __init__(self):
            pass

        @staticmethod
        def replace_in_paragraph(paragraph, buscar, remplazar, mantener_formato=False):
            paragraph_text = "".join([run.text for run in paragraph.runs])
            if buscar in paragraph_text:
                new_paragraph_text = paragraph_text.replace(buscar, remplazar)
                
                # Capturar el formato original si mantener_formato es True
                formato_original = None
                if mantener_formato and paragraph.runs:
                    formato_original = paragraph.runs[0].font
                
                paragraph.clear()
                new_run = paragraph.add_run(new_paragraph_text)
                
                # Aplicar el formato original si mantener_formato es True
                if mantener_formato and formato_original:
                    new_run.font.size = formato_original.size
                    new_run.font.name = formato_original.name
                    new_run.font.bold = formato_original.bold
                    new_run.font.italic = formato_original.italic
                    new_run.font.underline = formato_original.underline
                    new_run.font.color.rgb = formato_original.color.rgb
                   
                    
                return True
            return False


    def style_text(text, size, bold, ital, under, font_name):

        font = text.font
        font.size = size
        font.bold = bold
        font.italic = ital
        font.underline = under
        font.name = font_name

    session = GetParams("session")
    module = GetParams("module")
    global document, officeWord_session
        
    if not session:
        session = 'default'    
        
    try:
        if not officeWord_session : #type:ignore
            officeWord_session = {}
    except NameError:
        officeWord_session = {}

    if module == "new":
        officeWord_session[session] = Document()

    if module == "open":
        path = GetParams("path")

        officeWord_session[session] = Document(path)

    if module == "read":

        result = GetParams("result")

        read_path = os.path.join(base_path, "tmp")
        try:
            os.mkdir(read_path)
        except:
            pass
        print(read_path)

        officeWord_session[session].save(os.path.join(read_path, "tmp.docx")) #create temporal file
        text = docx2txt.process(os.path.join(read_path, "tmp.docx"))
        os.unlink(os.path.join(read_path, "tmp.docx")) #delete file

        if result:
            SetVar(result, text)
    
    if module == "readTable":

        result = GetParams("result")
        tableDoc = []
        for table in officeWord_session[session].tables:
            table_ = []
            for row in table.rows:
                array_row = []
                for cell in row.cells:
                    if len(array_row) > 0:
                        if array_row[-1] != cell.text:
                            array_row.append(cell.text)
                    else:
                        array_row.append(cell.text)
                table_.append(array_row)
            tableDoc.append(table_)
        if result:
            SetVar(result, tableDoc)

    if module == "addDataTable": 
        numTable = int(GetParams("numTable")) - 1
        data = GetParams("data")
        
        try:
            table = officeWord_session[session].tables[numTable]
            
            if data:
                data = eval(data)
                for i in range(len(data)):
                    for j in range(len(data[0])):
                       table.cell(i, j).text = data[i][j]
                        
            else:
                raise Exception("No data provided")
        
        except Exception as e:
            print("\x1B[" + "31;40mError\x1B[" + "0m")
            PrintException()
            raise e
    
    if module == "addTextBookmark":

        bookmark_searched = GetParams("bookmark")
        text = GetParams("text")
        clean = GetParams("Clean")

        try:
            tmp_doc = Document()
            # Generate content in tmp_doc document
            tmp_doc.add_paragraph(text)
            # Reference the tmp_doc XML content
            tmp_doc_body = tmp_doc._element.body


            ele = officeWord_session[session]._element[0]
            bookmarks_list = ele.findall('.//' + qn('w:bookmarkStart'))
            for bookmark in bookmarks_list:
                # print(bookmark)
                name = bookmark.get(qn('w:name'))
                if name == bookmark_searched:
                    par = bookmark.getparent()

                    if clean:
                        next_element = bookmark.getnext()
                        if not isinstance(next_element, docx.oxml.CT_R):
                            next_element = next_element.getnext()
                        t = next_element.findall('.//' + qn('w:t'))
                        if len(t) == 1:
                            t[0].text = text
                    elif isinstance(par, docx.oxml.CT_P):
                        bookmark_par_parent = par.getparent()
                        index = bookmark_par_parent.index(par)
                        for child in tmp_doc_body:
                            bookmark_par_parent.insert(index, child)
                            index = index + 1

                    break
                else:
                    name = None

            if not name:
                raise Exception("Bookmark not found")

        except Exception as e:
            PrintException()
            raise e

    if module == "save":

        path = GetParams("path")

        if path:
            if not path.endswith(".docx"):
                path += ".docx"
            officeWord_session[session].save(path)

    if module == "write":

        text = GetParams("text")
        type_ = GetParams("type")
        align = GetParams("align")
        size = GetParams("size")
        bold = GetParams("bold")
        ital = GetParams("italic")
        under = GetParams("underline")
        font_name = GetParams("font_name")

        if size:
            size = Pt(int(size))
        if bold:
            bold = eval(bold)
        if ital:
            ital = eval(ital)
        if under:
            under = eval(under)

        if align == "left" or None:
            align = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "center":
            align = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            align = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "justify":
            align = WD_ALIGN_PARAGRAPH.JUSTIFY

        if type_ == "title":
            print("title")
            t = officeWord_session[session].add_heading(level = 0)
            run = t.add_run(text)
            style_text(run, size, bold, ital, under, font_name)
            t.alignment = align
        elif type_ == "h1":
            t = officeWord_session[session].add_heading(level =1)
            run = t.add_run(text)
            style_text(run, size, bold, ital, under, font_name)
            t.alignment = align
        elif type_ == "h2":
            t = officeWord_session[session].add_heading(level = 2)
            run = t.add_run(text)
            style_text(run, size, bold, ital, under, font_name)
            t.alignment = align
        elif type_ == "p":
            texto = text.split("\\n ")
            for line in texto:
                t = officeWord_session[session].add_paragraph()
                run = t.add_run(line)
                style_text(run, size, bold, ital, under, font_name)
                t.alignment = align
        elif type_ == "bp":
            texto = text.split("\\n ")
            for line in texto:
                t = officeWord_session[session].add_paragraph(style='List Bullet')
                run = t.add_run(line)
                style_text(run, size, bold, ital, under, font_name)
                t.alignment = align
        elif type_ == "ln":
            texto = text.split("\\n ")
            for line in texto:
                t = officeWord_session[session].add_paragraph(style='List Number')
                run = t.add_run(line)
                style_text(run, size, bold, ital, under, font_name)
                t.alignment = align
        else:
            raise Exception("No se ha seleccionado tipo de texto")

    if module == "close":
        officeWord_session.pop(session)

    if module == "new_page":
        officeWord_session[session].add_page_break()

    if module == "add_pic":

        img_path = GetParams("img_path")
        width = GetParams("width")
        height = GetParams("height")
        if width == "" or width == None:
            width = None
        else:
            width = Mm(int(width))
            
        if height == "" or height == None:
            height = None
        else:
            height = Mm(int(height))


        officeWord_session[session].add_picture(img_path, width=width, height=height)

    if module == "to_pdf":
        try:

            from_ = GetParams("from")
            to_ = GetParams("to")

            if not to_.endswith(".pdf"):
                to_ += ".pdf"

            from_ = os.path.normpath(from_)
            to_ = os.path.normpath(to_)

            options = ' -f "' + from_ + '" -O "' + to_ + '" -T wdFormatPDF'
            options = ['-f', from_, "-O", to_, "-T", "wdFormatPDF"]

            run_ = [docto] + options
            print("\n\n",run_)
            con = Popen(run_, shell=True, stdout=PIPE, stderr=PIPE)
            a = con.communicate()
        except Exception as e:
            PrintException()
            raise e

    ## Modificado por Mijahil Franchi: Ubica el parrafo en el que se encuentra un texto
    if module == "search_text":
        parrafos = officeWord_session[session].paragraphs
        text_buscar = GetParams("text_search")
        variable = GetParams("variable")
        posicion = 0
        posiciones = list()
        try:
            for parrafo in parrafos:
                if text_buscar in parrafo.text:
                    posiciones.append(posicion)
                posicion = posicion+1
        except Exception as e:
            PrintException()
            raise e

        if posiciones:
            SetVar(variable, posiciones)

    ## Modificado por Mijahil Franchi: Cuenta los parrafos de un documento
    if module == "count_paragraphs":
        variable = GetParams("variable")
        
        try:
            parrafos = officeWord_session[session].paragraphs
            cantidad = len(parrafos)
        except Exception as e:
            PrintException()
            raise e

        if cantidad:
            SetVar(variable, cantidad)
            
    if module == "getParagraphs":
        result = GetParams("result")
        
        try:
            parrafos = officeWord_session[session].paragraphs
            parrafos_ = {}
            for i, parrafo in enumerate(parrafos):
                parrafos_["Paragraph "+str(i+1)] = parrafo.text
            
            SetVar(result, parrafos_)
        except Exception as e:
            SetVar(result, False)
            PrintException()
            raise e
        
    if module == "addParagraph":
        number = GetParams("number")
        text = GetParams("text")
        align = GetParams("align")
        size = GetParams("size")
        bold = GetParams("bold")
        ital = GetParams("italic")
        under = GetParams("underline")
        font_name = GetParams("font_name")
       
        try:
            number = int(number)-1
            parrafo = officeWord_session[session].paragraphs[number]
            parrafo.insert_paragraph_before(text)

            if size:
                size = Pt(int(size))
            if bold:
                bold = eval(bold)
            if ital:
                ital = eval(ital)
            if under:
                under = eval(under)

            if align == "left" or None:
                align = WD_ALIGN_PARAGRAPH.LEFT
            elif align == "center":
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                align = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == "justify":
                align = WD_ALIGN_PARAGRAPH.JUSTIFY

            for run in parrafo.runs:
                style_text(run, size, bold, ital, under, font_name)
            parrafo.alignment = align
    
        except Exception as e:
            PrintException()
            raise e
    
    if module == "addRun":
        number = GetParams("number")
        text = GetParams("text")
        align = GetParams("align")
        size = GetParams("size")
        bold = GetParams("bold")
        ital = GetParams("italic")
        under = GetParams("underline")
        font_name = GetParams("font_name")
       
        try:
            number = int(number)-1
            parrafo = officeWord_session[session].paragraphs[number]
            run = parrafo.add_run(text)

            if size:
                size = Pt(int(size))
            if bold:
                bold = eval(bold)
            if ital:
                ital = eval(ital)
            if under:
                under = eval(under)

            if align == "left" or None:
                align = WD_ALIGN_PARAGRAPH.LEFT
            elif align == "center":
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                align = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == "justify":
                align = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            style_text(run, size, bold, ital, under, font_name)
            parrafo.alignment = align
    
        except Exception as e:
            PrintException()
            raise e

    if module == "clearParagraph":
        number = GetParams("number")
        result = GetParams("result")
        
        try:
            number = int(number)-1
            parrafo = officeWord_session[session].paragraphs[number]
            clear = parrafo.clear()
            
            SetVar(result, True)
        except Exception as e:
            SetVar(result, False)
            PrintException()
            raise e
    
    ## Modificado por Mijahil Franchi: Busca y remplaza el contenido de un texto
    if module == "search_replace_text":
        
        variable = GetParams("variable")
        parrafos = GetParams("parrafos")
        buscar = GetParams("text_search")
        remplazar = GetParams("text_replace")
        mantener_formato = eval(GetParams("mantener_formato") or "False")
        result = False
        posicion = 0

        try:

            paragraphs = officeWord_session[session].paragraphs
            if parrafos:
                for line in parrafos.split(','):
                    paragraph = paragraphs[int(line)]
                    result = DocxModule.replace_in_paragraph(paragraph, buscar, remplazar, mantener_formato)
            else:
                for paragraph in paragraphs:
                    result = DocxModule.replace_in_paragraph(paragraph, buscar, remplazar, mantener_formato)

            SetVar(variable, result)
                
        except Exception as e:
            SetVar(variable, result)
            PrintException()
            raise e
except Exception as e:
    print(traceback.print_exc())
    PrintException()
    raise e
    
    